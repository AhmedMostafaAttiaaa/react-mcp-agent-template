"""File parsing, chunking, embedding, and index-building for the RAG pipeline.
Everything here runs on CPU: parsing is pure-Python/pandas, embeddings go through
a local Ollama model, and indexing writes to disk (see rag_store.py).
"""
import io
import json
from dataclasses import dataclass

import httpx
import ollama
import pandas as pd
from docx import Document
from pypdf import PdfReader

# See agent/providers/ollama_provider.py — httpx's flat 5s default is too tight
# once OLLAMA_HOST is a real network host rather than localhost.
_DEFAULT_TIMEOUT = httpx.Timeout(120.0, connect=30.0)


@dataclass
class Chunk:
    text: str
    source: str
    chunk_index: int


def extract_text(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    if lower.endswith(".csv"):
        df = pd.read_csv(io.BytesIO(data))
        return df.to_csv(index=False)
    if lower.endswith(".xlsx"):
        df = pd.read_excel(io.BytesIO(data), engine="openpyxl")
        return df.to_csv(index=False)
    if lower.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if lower.endswith(".docx"):
        document = Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)
    if lower.endswith(".doc"):
        try:
            document = Document(io.BytesIO(data))
            return "\n".join(p.text for p in document.paragraphs)
        except Exception as exc:
            raise ValueError(
                "Legacy .doc files aren't reliably parseable without extra dependencies. "
                "Please re-save the file as .docx and re-upload."
            ) from exc
    raise ValueError(f"Unsupported file type: {filename}")


def chunk_text(text: str, source: str, chunk_size: int, chunk_overlap: int) -> list:
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size")

    words = text.split()
    if not words:
        return []

    chunks = []
    step = chunk_size - chunk_overlap
    index = 0
    for start in range(0, len(words), step):
        window = words[start : start + chunk_size]
        if not window:
            break
        chunks.append(Chunk(text=" ".join(window), source=source, chunk_index=index))
        index += 1
        if start + chunk_size >= len(words):
            break
    return chunks


def make_ollama_embedder(model: str, host: str):
    client = ollama.Client(host=host, timeout=_DEFAULT_TIMEOUT)

    def embed(text: str) -> list:
        response = client.embeddings(model=model, prompt=text)
        return response["embedding"]

    return embed


def embed_chunks(chunks: list, embed_fn) -> list:
    return [embed_fn(chunk.text) for chunk in chunks]


def build_vector_index(chunks: list, embed_fn, store) -> int:
    if not chunks:
        return 0
    embeddings = embed_chunks(chunks, embed_fn)
    ids = [f"{c.source}-{c.chunk_index}" for c in chunks]
    documents = [c.text for c in chunks]
    metadatas = [{"source": c.source, "chunk_index": c.chunk_index} for c in chunks]
    store.add(ids=ids, embeddings=embeddings, documents=documents, metadatas=metadatas)
    return len(chunks)


_EXTRACTION_PROMPT = """Extract entities and relationships from the text below as a JSON array of
objects with keys "subject", "relation", "object". Use short, normalized entity names. Return
ONLY the JSON array, with no extra commentary.

Text:
{text}

JSON:"""


def extract_triples(chunk_text_value: str, provider) -> list:
    """Uses the currently active generation provider (ollama or groq) for entity/
    relationship extraction — no separate NER dependency needed. Fails soft: any
    parsing problem just yields zero triples for that chunk instead of raising.
    """
    prompt = _EXTRACTION_PROMPT.format(text=chunk_text_value)
    result = provider.generate(prompt, stop=None)
    raw = result.text.strip()
    start = raw.find("[")
    end = raw.rfind("]")
    if start == -1 or end == -1:
        return []
    try:
        triples = json.loads(raw[start : end + 1])
    except json.JSONDecodeError:
        return []
    return [t for t in triples if isinstance(t, dict) and {"subject", "relation", "object"} <= t.keys()]


def build_graph_index(chunks: list, provider, store) -> int:
    triple_count = 0
    for chunk in chunks:
        for triple in extract_triples(chunk.text, provider):
            store.add_triple(triple["subject"], triple["relation"], triple["object"], source=chunk.source)
            triple_count += 1
    store.save()
    return triple_count
